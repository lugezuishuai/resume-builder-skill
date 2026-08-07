#!/usr/bin/env python3
"""
Professional tech-resume generator.

Reads a single JSON resume file and produces HTML, DOCX, and PDF outputs using
a fixed navy-blue consulting-style template. The visual system (colors, spacing,
typography, banner styles) mirrors the style refined in the author's own resume
and is intentionally not parameterised — keeping the template deterministic
produces consistently polished output.

Usage:
    python generate_resume.py --data resume.json [--outdir ./output] [--photo photo.jpg]

Outputs (in --outdir, default current dir):
    resume.html   — single-file HTML (inline CSS), WeasyPrint-compatible
    resume.docx   — editable Word document
    resume.pdf    — rendered via WeasyPrint from the HTML (preferred output)

JSON schema — see references/data-schema.md for full spec.
"""
from __future__ import annotations

import argparse
from html import escape
from importlib.util import find_spec
import json
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

# ---------- Color palette (consulting / navy-blue) ----------
NAVY = "#1e40af"
NAVY_RGB = (0x1E, 0x40, 0xAF)
DARK = "#111827"
DARK_RGB = (0x11, 0x18, 0x27)
GRAY = "#4b5563"
GRAY_RGB = (0x4B, 0x55, 0x63)
BODY = "#1f2937"
BODY_RGB = (0x1F, 0x29, 0x37)
LIGHT = "#6b7280"
WHITE = "#ffffff"
BANNER_BG = "#eff6ff"
BORDER_BLUE = "#1e40af"
BOOTSTRAP_ENV = "RESUME_BUILDER_BOOTSTRAPPED"
SYSTEM_DEPS_ENV = "RESUME_BUILDER_SYSTEM_DEPS_BOOTSTRAPPED"
PACKAGE_IMPORTS = {
    "python-docx": "docx",
    "weasyprint": "weasyprint",
}


def _runtime_venv_dir() -> Path:
    """Return the isolated environment used for this skill's Python packages."""
    configured = os.environ.get("RESUME_BUILDER_VENV")
    if configured:
        return Path(configured).expanduser()
    if os.name == "nt":
        cache_root = Path(os.environ.get("LOCALAPPDATA", Path.home() / "AppData" / "Local"))
    else:
        cache_root = Path(os.environ.get("XDG_CACHE_HOME", Path.home() / ".cache"))
    return cache_root / "resume-builder" / "venv"


def _venv_python(venv_dir: Path) -> Path:
    return venv_dir / ("Scripts/python.exe" if os.name == "nt" else "bin/python")


def _requested_formats(args: argparse.Namespace) -> set[str]:
    """Resolve the requested artifacts while retaining the legacy flags."""
    if args.format != "all":
        return {args.format}
    if args.html_only:
        return {"html"}
    if args.no_pdf:
        return {"html", "docx"}
    return {"html", "docx", "pdf"}


def _required_packages(formats: set[str]) -> list[str]:
    """Return only the Python packages required by the requested artifacts."""
    packages = []
    if "docx" in formats:
        packages.append("python-docx")
    if "pdf" in formats:
        packages.append("weasyprint")
    return packages


def ensure_runtime_dependencies(formats: set[str]) -> None:
    """Install missing optional packages in an isolated venv, then re-run the script.

    `pip` cannot safely modify Homebrew- or OS-managed Python installations on
    many machines. Keeping the runtime in a per-user cache makes first use
    self-contained and leaves the caller's Python environment untouched.
    """
    required = _required_packages(formats)
    missing = [package for package in required if find_spec(PACKAGE_IMPORTS[package]) is None]
    if not missing:
        return

    if os.environ.get(BOOTSTRAP_ENV) == "1":
        missing_list = ", ".join(missing)
        raise RuntimeError(f"自动安装后仍缺少依赖：{missing_list}")

    venv_dir = _runtime_venv_dir()
    venv_python = _venv_python(venv_dir)
    try:
        if not venv_python.is_file():
            print(f"正在创建 Skill 运行环境：{venv_dir}")
            subprocess.run([sys.executable, "-m", "venv", str(venv_dir)], check=True)
        print(f"正在安装运行所需依赖：{', '.join(required)}")
        subprocess.run(
            [str(venv_python), "-m", "pip", "install", *required],
            check=True,
        )
    except (OSError, subprocess.CalledProcessError) as exc:
        raise RuntimeError(
            "无法自动安装 Python 依赖。请确认网络可用、当前 Python 包含 venv 模块，"
            "且用户缓存目录可写。"
        ) from exc

    env = os.environ.copy()
    env[BOOTSTRAP_ENV] = "1"
    command = [str(venv_python), str(Path(__file__).resolve()), *sys.argv[1:]]
    raise SystemExit(subprocess.run(command, env=env).returncode)


def _configure_pdf_environment() -> None:
    """Provide a writable font cache and Homebrew libraries to WeasyPrint."""
    font_cache = Path(tempfile.gettempdir()) / "resume-builder-fontconfig"
    font_cache.mkdir(parents=True, exist_ok=True)
    os.environ["XDG_CACHE_HOME"] = str(font_cache)

    if sys.platform == "darwin":
        brew = shutil.which("brew")
        if brew:
            prefix = subprocess.run(
                [brew, "--prefix"], capture_output=True, check=False, text=True
            ).stdout.strip()
            lib_dir = Path(prefix) / "lib"
            if lib_dir.is_dir():
                existing = os.environ.get("DYLD_FALLBACK_LIBRARY_PATH", "")
                os.environ["DYLD_FALLBACK_LIBRARY_PATH"] = (
                    f"{lib_dir}:{existing}" if existing else str(lib_dir)
                )


def _install_macos_pdf_libraries() -> None:
    """Install the native Pango stack required by WeasyPrint on macOS."""
    brew = shutil.which("brew")
    if not brew:
        raise RuntimeError(
            "PDF 渲染缺少 Pango/GObject，且未找到 Homebrew。请安装 Homebrew 后重试，"
            "或在已具备 Pango 的环境中运行。"
        )
    print("正在安装 WeasyPrint 所需的 macOS 系统库：pango")
    try:
        subprocess.run([brew, "install", "pango"], check=True)
    except (OSError, subprocess.CalledProcessError) as exc:
        raise RuntimeError("无法自动安装 Pango；请确认 Homebrew 可用并允许安装系统依赖。") from exc
    _configure_pdf_environment()


def _restart_after_system_dependency_install() -> None:
    env = os.environ.copy()
    env[SYSTEM_DEPS_ENV] = "1"
    command = [sys.executable, str(Path(__file__).resolve()), *sys.argv[1:]]
    raise SystemExit(subprocess.run(command, env=env).returncode)


def _write_pdf(html_path: Path, pdf_path: Path) -> None:
    """Render a PDF and recover once from a missing native macOS library."""
    _configure_pdf_environment()
    try:
        from weasyprint import HTML as WPHTML
        WPHTML(filename=str(html_path), base_url=str(html_path.parent)).write_pdf(str(pdf_path))
    except OSError as exc:
        if sys.platform == "darwin" and os.environ.get(SYSTEM_DEPS_ENV) != "1":
            _install_macos_pdf_libraries()
            _restart_after_system_dependency_install()
        raise RuntimeError(
            "WeasyPrint 无法加载系统图形库。Linux 请安装 Pango/Cairo，Windows 请配置 "
            "MSYS2 Pango 或使用 WSL；macOS 请确认 Homebrew 的 pango 可用。"
        ) from exc

# ============================================================
#                        HTML GENERATION
# ============================================================

HTML_CSS = """
  @page { size: A4; margin: 10mm 12mm 10mm 12mm; }
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body {
    font-family: -apple-system, "PingFang SC", "Microsoft YaHei", "Helvetica Neue",
                 "Noto Sans CJK SC", Arial, sans-serif;
    font-size: 9pt;
    line-height: 1.45;
    color: #1f2937;
  }
  /* ---------- Header ---------- */
  .header { position: relative; margin-bottom: 4px; padding-right: 95px; min-height: 105px; }
  .photo { position: absolute; top: 0; right: 0; width: 80px; height: 108px;
           object-fit: cover; border: 1px solid #d1d5db; }
  .name { font-size: 24pt; font-weight: 700; color: #1e40af; letter-spacing: 3px;
          line-height: 1.1; margin-bottom: 5px; }
  .title-tag { display: inline-block; background: #1e40af; color: #fff;
               font-size: 9.5pt; font-weight: 600; padding: 2px 9px; border-radius: 3px;
               margin-bottom: 6px; letter-spacing: 1px; }
  .contact { font-size: 8.5pt; color: #4b5563; border-bottom: 2px solid #1e40af;
             padding-bottom: 4px; display: flex; flex-wrap: wrap; gap: 3px 16px; }
  .contact a { color: #4b5563; text-decoration: none; }
  /* ---------- Section ---------- */
  .section { margin-bottom: 4px; }
  .section-title { font-size: 12pt; font-weight: 700; color: #1e40af;
                   border-bottom: 1.5px solid #1e40af; padding-bottom: 1px;
                   margin-bottom: 3px; letter-spacing: 1px; }
  .summary-text { font-size: 9.2pt; line-height: 1.5; text-align: justify; margin: 0; }
  .skill-row { margin-bottom: 0; font-size: 8.8pt; line-height: 1.45; }
  /* ---------- Work / Projects ---------- */
  .company-header { display: flex; justify-content: space-between; align-items: baseline;
                    margin-bottom: 2px; }
  .company-name { font-size: 11pt; font-weight: 700; color: #111827; }
  .company-meta { font-size: 8.5pt; color: #4b5563; }
  .company-intro { font-size: 9pt; color: #1f2937; margin-bottom: 4px;
                   line-height: 1.5; text-align: justify; }
  .project-title { position: relative; margin: 5px 0 2px 0; padding: 3px 12px;
                   background: #1e40af; font-size: 10pt; font-weight: 700; color: #ffffff;
                   letter-spacing: 0.5px; line-height: 1.35; border-radius: 2px;
                   page-break-after: avoid; }
  .project-intro { font-size: 8.8pt; line-height: 1.5; color: #1f2937; margin-bottom: 1px;
                   text-align: justify; }
  .sub-title { position: relative; margin: 3px 0 1px 0; padding: 1px 10px 1px 13px;
               background: #eff6ff; font-size: 9pt; font-weight: 700; color: #1e40af;
               letter-spacing: 0.5px; line-height: 1.35; overflow: hidden;
               page-break-after: avoid; }
  .sub-title::before { content: ""; position: absolute; left: 0; top: 0; bottom: 0;
                       width: 4px; background: #1e40af; }
  .project-body { font-size: 8.8pt; line-height: 1.5; color: #1f2937; }
  .bullet { padding-left: 9px; position: relative; margin-bottom: 1px; text-align: justify; }
  .bullet::before { content: "·"; position: absolute; left: 0; color: #1e40af; font-weight: 700; }
  /* ---------- Education / Certs ---------- */
  .edu-row { display: flex; justify-content: space-between; align-items: baseline;
             font-size: 9.2pt; margin-bottom: 1px; }
  .edu-name { font-weight: 700; color: #111827; }
  .edu-meta { font-size: 8.5pt; color: #4b5563; }
  .cert-line { font-size: 9pt; color: #1f2937; }
  .cert-line .sep { color: #1e40af; margin: 0 6px; }
  strong { font-weight: 700; color: #111827; }
  .page-break { page-break-before: always; }
"""


def _escape_text(value: Any) -> str:
    """Escape user-provided text before inserting it into HTML."""
    return escape(str(value), quote=True)


def _inline_bold(text: Any) -> str:
    """Escape user text, then convert non-nested **bold** markers to HTML."""
    import re
    escaped = _escape_text(text)
    return re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)


def _safe_url(value: Any) -> str | None:
    """Return an absolute HTTP(S) URL, or None for unsafe/invalid links."""
    url = str(value)
    parsed = urlparse(url)
    if parsed.scheme.lower() in {"http", "https"} and parsed.netloc:
        return url
    return None


def _contact_line(c: dict) -> str:
    parts = []
    if c.get("age"):
        parts.append(f"年龄：{_escape_text(c['age'])}")
    if c.get("phone"):
        parts.append(f"电话：{_escape_text(c['phone'])}")
    if c.get("email"):
        parts.append(f"邮箱：{_escape_text(c['email'])}")
    if c.get("wechat"):
        parts.append(f"微信：{_escape_text(c['wechat'])}")
    if c.get("blog"):
        raw_url = c["blog"]
        url = _safe_url(raw_url)
        if url:
            label = url.replace("https://", "").replace("http://", "").rstrip("/")
            parts.append(f'博客：<a href="{_escape_text(url)}">{_escape_text(label)}</a>')
        else:
            parts.append(f"博客：{_escape_text(raw_url)}")
    extra = c.get("extra") or []
    for item in extra:
        if isinstance(item, dict):
            label, raw_url = item.get("label",""), item.get("url","")
            url = _safe_url(raw_url)
            if url:
                visible_url = url.replace("https://", "").replace("http://", "").rstrip("/")
                parts.append(f'{_escape_text(label)}：<a href="{_escape_text(url)}">{_escape_text(visible_url)}</a>')
            else:
                parts.append(_escape_text(label))
        else:
            parts.append(_escape_text(item))
    return "<span>" + "</span><span>".join(parts) + "</span>"


def build_html(data: dict, photo_path: str | None = None) -> str:
    name = _escape_text(data.get("name", "姓名"))
    title = _escape_text(data.get("title", "职位"))
    contact = data.get("contact", {})
    summary = data.get("summary", "")
    skills = data.get("skills", [])
    experience = data.get("experience", [])
    education = data.get("education", [])
    certs = data.get("certs", [])

    photo_html = ""
    if photo_path and os.path.isfile(photo_path):
        photo_html = f'<img class="photo" src="file://{_escape_text(os.path.abspath(photo_path))}" alt="证件照">'

    # Skills
    skill_html = []
    for s in skills:
        label = s.get("label", "")
        desc = s.get("desc", "")
        skill_html.append(
            f'<div class="skill-row"><strong>{_escape_text(label)}</strong>{_inline_bold(desc)}</div>'
        )
    skills_block = "\n  ".join(skill_html)

    # Experience
    exp_parts = []
    for idx, company in enumerate(experience):
        c_name = _inline_bold(company.get("company", ""))
        c_role = _inline_bold(company.get("role", ""))
        c_meta = _inline_bold(company.get("meta", ""))
        c_intro = _inline_bold(company.get("intro", ""))
        page_break_cls = ' style="margin-top:10px;"' if idx > 0 else ""

        exp_parts.append(f"""  <div class="company-header"{page_break_cls}>
    <div class="company-name">{c_name}<span style="font-weight:400;color:#4b5563;font-size:10pt;">｜{c_role}</span></div>
    <div class="company-meta">{c_meta}</div>
  </div>
  <div class="company-intro">{c_intro}</div>""")

        for proj in company.get("projects", []):
            pb_cls = ' class="project-title page-break"' if proj.get("page_break") else ' class="project-title"'
            exp_parts.append(f'  <div{pb_cls}>{_inline_bold(proj.get("title", ""))}</div>')
            if proj.get("intro"):
                exp_parts.append(
                    f'  <div class="project-intro">{_inline_bold(proj["intro"])}</div>'
                )
            for sec in proj.get("sections", []):
                st = sec.get("subtitle")
                if st:
                    exp_parts.append(f'  <div class="sub-title">{_inline_bold(st)}</div>')
                exp_parts.append('  <div class="project-body">')
                for b in sec.get("bullets", []):
                    exp_parts.append(
                        f'    <div class="bullet">{_inline_bold(b)}</div>'
                    )
                exp_parts.append('  </div>')
            # Allow raw bullets directly under project (no subsection)
            bullets = proj.get("bullets", [])
            if bullets:
                exp_parts.append('  <div class="project-body">')
                for b in bullets:
                    exp_parts.append(
                        f'    <div class="bullet">{_inline_bold(b)}</div>'
                    )
                exp_parts.append('  </div>')
    exp_block = "\n".join(exp_parts)

    # Education
    edu_html = []
    for e in education:
        edu_html.append(
            f'''<div class="edu-row">
    <div class="edu-name">{_inline_bold(e.get("school", ""))}<span style="font-weight:400;color:#374151;">｜{_inline_bold(e.get("degree", ""))}</span></div>
    <div class="edu-meta">{_inline_bold(e.get("meta", ""))}</div>
  </div>'''
        )
    edu_block = "\n  ".join(edu_html)

    # Certifications
    cert_parts = []
    for i, c in enumerate(certs):
        if isinstance(c, dict):
            text = c.get("text", "")
            bold = c.get("bold", False)
        else:
            text = str(c); bold = (i == 0)
        rendered_text = _inline_bold(text)
        cert_parts.append(f"<strong>{rendered_text}</strong>" if bold else rendered_text)
    cert_block = ('<span class="sep">·</span>'.join(cert_parts))

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>{name} - 简历</title>
<style>{HTML_CSS}</style>
</head>
<body>

<div class="header">
  {photo_html}
  <div class="name">{name}</div>
  <div class="title-tag">{title}</div>
  <div class="contact">
    {_contact_line(contact)}
  </div>
</div>

<div class="section">
  <div class="section-title">个人概述</div>
  <div class="summary-text">{_inline_bold(summary)}</div>
</div>

<div class="section">
  <div class="section-title">技术栈 / 核心能力</div>
  {skills_block}
</div>

<div class="section">
  <div class="section-title">工作经历</div>
{exp_block}
</div>

<div class="section">
  <div class="section-title">教育经历</div>
  {edu_block}
</div>

<div class="section">
  <div class="section-title">证书 / 荣誉</div>
  <div class="cert-line">{cert_block}</div>
</div>

</body>
</html>"""
    return html


# ============================================================
#                        DOCX GENERATION
# ============================================================

def build_docx(data: dict, out_path: str, photo_path: str | None = None):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    FONT = "PingFang SC"

    def set_font(run, size=9, bold=False, color_rgb=BODY_RGB, font=FONT):
        run.font.name = font
        run.font.size = Pt(size)
        run.font.bold = bold
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
        for k in ("w:eastAsia", "w:ascii", "w:hAnsi"):
            rFonts.set(qn(k), font)

    def shade(paragraph, color_hex):
        pPr = paragraph._element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
        shd.set(qn("w:fill"), color_hex)
        pPr.append(shd)

    def add_border(paragraph, **kwargs):
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        for side, (color, sz, space) in kwargs.items():
            b_el = OxmlElement(f"w:{side}")
            b_el.set(qn("w:val"),"single"); b_el.set(qn("w:sz"),str(sz))
            b_el.set(qn("w:space"),str(space)); b_el.set(qn("w:color"),color)
            pBdr.append(b_el)
        pPr.append(pBdr)

    def add_runs_with_bold(p, text, size=9, color=BODY_RGB, ls=1.45):
        """Parse **bold** segments and add runs."""
        import re
        parts = re.split(r"(\*\*.+?\*\*)", text)
        pf = p.paragraph_format
        pf.line_spacing = ls
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                r = p.add_run(part[2:-2]); set_font(r, size=size, bold=True, color_rgb=DARK_RGB)
            else:
                r = p.add_run(part); set_font(r, size=size, color_rgb=color)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    sec.top_margin=Cm(1.0); sec.bottom_margin=Cm(1.0)
    sec.left_margin=Cm(1.2); sec.right_margin=Cm(1.2)

    # Header
    table = doc.add_table(rows=1, cols=2); table.autofit = False
    for cell, w in zip(table.rows[0].cells, [Cm(15.7), Cm(2.6)]):
        cell.width = w
    tbl = table._tbl; tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        b_el = OxmlElement(f"w:{edge}")
        b_el.set(qn("w:val"),"nil"); tblBorders.append(b_el)
    tblPr.append(tblBorders)
    left_cell = table.rows[0].cells[0]; right_cell = table.rows[0].cells[1]
    tcPr = right_cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign"); vAlign.set(qn("w:val"),"top"); tcPr.append(vAlign)

    p = left_cell.paragraphs[0]
    pf = p.paragraph_format; pf.space_after=Pt(3); pf.line_spacing=1.0
    r = p.add_run(data.get("name","")); set_font(r,size=24,bold=True,color_rgb=NAVY_RGB)

    p = left_cell.add_paragraph()
    pf = p.paragraph_format; pf.space_after=Pt(5); pf.line_spacing=1.0
    shade(p, "1E40AF")
    r = p.add_run("  " + data.get("title","") + "  ")
    set_font(r,size=9.5,bold=True,color_rgb=(0xFF,0xFF,0xFF))

    # Contact
    contact = data.get("contact", {})
    p = left_cell.add_paragraph()
    pf = p.paragraph_format; pf.space_after=Pt(0); pf.line_spacing=1.3
    add_border(p, bottom=("1E40AF","12","1"))
    pieces=[]
    if contact.get("age"): pieces.append(f"年龄：{contact['age']}")
    if contact.get("phone"): pieces.append(f"电话：{contact['phone']}")
    if contact.get("email"): pieces.append(f"邮箱：{contact['email']}")
    if contact.get("wechat"): pieces.append(f"微信：{contact['wechat']}")
    contact_text = "    ".join(pieces) + ("    " if pieces else "")
    r = p.add_run(contact_text); set_font(r,size=8.5,color_rgb=GRAY_RGB)
    if contact.get("blog"):
        r = p.add_run("博客："); set_font(r,size=8.5,color_rgb=GRAY_RGB)
        url = contact["blog"]; label = url.replace("https://","").replace("http://","").rstrip("/")
        # Add hyperlink
        part = p.part
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        r_id = part.relate_to(url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True)
        hyperlink = OxmlElement("w:hyperlink"); hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        for k in ("w:eastAsia","w:ascii","w:hAnsi"): rFonts.set(qn(k),FONT)
        rPr.append(rFonts)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"),"17"); rPr.append(sz)
        c = OxmlElement("w:color"); c.set(qn("w:val"),"4B5563"); rPr.append(c)
        u = OxmlElement("w:u"); u.set(qn("w:val"),"none"); rPr.append(u)
        new_run.append(rPr)
        t = OxmlElement("w:t"); t.text = label; new_run.append(t)
        hyperlink.append(new_run); p._p.append(hyperlink)

    # Photo
    if photo_path and os.path.isfile(photo_path):
        p = right_cell.paragraphs[0]; pf = p.paragraph_format
        pf.space_before=Pt(0); pf.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(); r.add_picture(photo_path, width=Cm(2.2))

    p = doc.add_paragraph(); p.paragraph_format.space_after=Pt(1)

    # Section title helper
    def section_title(text):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before=Pt(3); pf.space_after=Pt(1); pf.line_spacing=1.25
        add_border(p, bottom=("1E40AF","8","1"))
        r = p.add_run(text); set_font(r,size=12,bold=True,color_rgb=NAVY_RGB)

    def banner(text, bg="1E40AF", fg=(0xFF,0xFF,0xFF), size=10, left_border=False):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before=Pt(5); pf.space_after=Pt(2); pf.line_spacing=1.3
        if left_border:
            add_border(p, left=("1E40AF","24","2"))
            pf.left_indent = Cm(0.22)
            shade(p, "EFF6FF")
            r = p.add_run(text); set_font(r,size=9,bold=True,color_rgb=NAVY_RGB)
        else:
            shade(p, bg)
            pf.left_indent = Cm(0.12)
            r = p.add_run("  " + text + "  "); set_font(r,size=size,bold=True,color_rgb=fg)

    def bullet(text, size=8.8, ls=1.45):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before=Pt(0); pf.space_after=Pt(0); pf.line_spacing=ls
        pf.left_indent=Cm(0.45); pf.first_line_indent=Cm(-0.25)
        r0 = p.add_run("· "); set_font(r0,size=size,bold=True,color_rgb=NAVY_RGB)
        add_runs_with_bold(p, text, size=size, color=BODY_RGB, ls=ls)
        # The first run was already the bullet; add_runs added runs after it but re-used p; fine

    def para_text(text, size=9, sb=0, sa=1, ls=1.5):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before=Pt(sb); pf.space_after=Pt(sa); pf.line_spacing=ls
        add_runs_with_bold(p, text, size=size, color=BODY_RGB, ls=ls)

    # Summary
    section_title("个人概述")
    para_text(data.get("summary",""), size=9.2, sa=1)

    # Skills
    section_title("技术栈 / 核心能力")
    for s in data.get("skills", []):
        label = s.get("label","")
        desc = s.get("desc","")
        p = doc.add_paragraph()
        pf = p.paragraph_format; pf.space_after=Pt(0); pf.line_spacing=1.45
        r = p.add_run(label); set_font(r,size=8.8,bold=True,color_rgb=DARK_RGB)
        r = p.add_run(desc); set_font(r,size=8.8,color_rgb=BODY_RGB)

    # Experience
    section_title("工作经历")
    for ci, company in enumerate(data.get("experience", [])):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before=Pt(1 if ci==0 else 6); pf.space_after=Pt(2); pf.line_spacing=1.25
        r = p.add_run(company.get("company","")); set_font(r,size=11,bold=True,color_rgb=DARK_RGB)
        r = p.add_run("｜" + company.get("role","")); set_font(r,size=9.5,color_rgb=GRAY_RGB)
        r = p.add_run("\t"); set_font(r,size=9)
        r = p.add_run(company.get("meta","")); set_font(r,size=8.5,color_rgb=GRAY_RGB)
        # Right tab
        pPr = p._element.get_or_add_pPr()
        tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab")
        tab.set(qn("w:val"),"right"); tab.set(qn("w:pos"),"9200"); tabs.append(tab); pPr.append(tabs)

        if company.get("intro"):
            para_text(company["intro"], size=9, sa=3)

        for proj in company.get("projects", []):
            if proj.get("page_break"):
                doc.add_page_break()
            banner(proj.get("title",""), bg="1E40AF")
            if proj.get("intro"):
                para_text(proj["intro"], size=8.8, sb=0, sa=1, ls=1.5)
            for sec in proj.get("sections", []):
                st = sec.get("subtitle")
                if st:
                    banner(st, left_border=True)
                for b in sec.get("bullets", []):
                    bullet(b)
            for b in proj.get("bullets", []):
                bullet(b)

    # Education
    section_title("教育经历")
    for e in data.get("education", []):
        p = doc.add_paragraph()
        pf = p.paragraph_format; pf.space_after=Pt(0); pf.line_spacing=1.4
        r = p.add_run(e.get("school","")); set_font(r,size=9.2,bold=True,color_rgb=DARK_RGB)
        r = p.add_run("｜" + e.get("degree","")); set_font(r,size=9.2,color_rgb=BODY_RGB)
        r = p.add_run("\t"); set_font(r,size=9)
        r = p.add_run(e.get("meta","")); set_font(r,size=8.5,color_rgb=GRAY_RGB)
        pPr = p._element.get_or_add_pPr()
        tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab")
        tab.set(qn("w:val"),"right"); tab.set(qn("w:pos"),"9200"); tabs.append(tab); pPr.append(tabs)

    # Certs
    section_title("证书 / 荣誉")
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.space_after=Pt(0)
    certs = data.get("certs", [])
    for i, c in enumerate(certs):
        if isinstance(c, dict):
            text=c.get("text",""); bold=c.get("bold", i==0)
        else:
            text=str(c); bold=(i==0)
        r = p.add_run(text); set_font(r,size=9,bold=bold,color_rgb=DARK_RGB if bold else BODY_RGB)
        if i < len(certs)-1:
            r = p.add_run("  ·  "); set_font(r,size=9,color_rgb=NAVY_RGB)

    doc.save(out_path)


# ============================================================
#                           MAIN
# ============================================================

def main():
    ap = argparse.ArgumentParser(description="Generate HTML/DOCX/PDF tech resume from JSON data.")
    ap.add_argument("--data", required=True, help="Path to resume JSON file")
    ap.add_argument("--outdir", default=".", help="Output directory (default: cwd)")
    ap.add_argument("--photo", default=None, help="Optional path to photo.jpg")
    ap.add_argument("--format", choices=["all", "html", "docx", "pdf"], default="all",
                    help="Artifact to keep (default: all)")
    ap.add_argument("--html-only", action="store_true", help="Legacy alias for --format html")
    ap.add_argument("--no-pdf", action="store_true", help="Legacy mode: generate HTML + DOCX")
    args = ap.parse_args()
    if args.format != "all" and (args.html_only or args.no_pdf):
        ap.error("--format cannot be combined with --html-only or --no-pdf")

    formats = _requested_formats(args)

    try:
        ensure_runtime_dependencies(formats)
    except RuntimeError as exc:
        ap.error(str(exc))

    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)

    outdir = Path(args.outdir); outdir.mkdir(parents=True, exist_ok=True)
    photo = args.photo

    html = build_html(data, photo) if formats & {"html", "pdf"} else ""

    if "html" in formats:
        html_path = outdir / "resume.html"
        html_path.write_text(html, encoding="utf-8")
        print(f"HTML → {html_path}")

    if "docx" in formats:
        docx_path = outdir / "resume.docx"
        try:
            build_docx(data, str(docx_path), photo)
            print(f"DOCX → {docx_path}")
        except ImportError as exc:
            ap.error(f"DOCX 依赖不可用：{exc}")

    if "pdf" in formats:
        pdf_path = outdir / "resume.pdf"
        if "html" in formats:
            _write_pdf(outdir / "resume.html", pdf_path)
        else:
            with tempfile.TemporaryDirectory(prefix="resume-builder-") as temp_dir:
                temp_html = Path(temp_dir) / "resume.html"
                temp_html.write_text(html, encoding="utf-8")
                _write_pdf(temp_html, pdf_path)
        print(f"PDF  → {pdf_path}")


if __name__ == "__main__":
    main()
